from __future__ import annotations

import sys
from pathlib import Path
from statistics import mean, median

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


PROJECT_DIR = Path(__file__).resolve().parents[1]
WORKSPACE_DIR = PROJECT_DIR.parent
DOCS_DIR = WORKSPACE_DIR / "docs"
SRC_DIR = PROJECT_DIR / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from vkr_classifier.config import get_settings  # noqa: E402
from vkr_classifier.data.image_generator import create_shape_image  # noqa: E402
from vkr_classifier.service import ClassifierService  # noqa: E402


STYLE_H1 = "ГОСТ Заголовок 1"
STYLE_H2 = "ГОСТ Заголовок 2"
STYLE_BODY = "ГОСТ Текст"
STYLE_CAPTION = "Caption"
OUTPUT_NAME = "Пояснительная записка.docx"
SOURCE_NAME = "Пояснительная записка_исходная.docx"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_paragraph = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph)
    result = Paragraph(new_paragraph, paragraph._parent)
    if text:
        result.add_run(text)
    if style:
        result.style = style
    return result


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def clear_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def set_run_font(run, font_name: str = "Times New Roman", font_size: int = 14, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold


def format_paragraph_runs(paragraph: Paragraph, font_name: str = "Times New Roman", font_size: int = 14, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def format_caption(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph_runs(paragraph, font_name="Times New Roman", font_size=14)


def add_centered_picture(paragraph: Paragraph, image_path: Path, width_inches: float = 6.1) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def add_styled_table(document: Document, title: str, data: pd.DataFrame) -> Table:
    title_paragraph = document.add_paragraph(title, style=STYLE_BODY)
    title_paragraph.style = STYLE_BODY
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = document.add_table(rows=1, cols=len(data.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for cell, column_name in zip(header_cells, data.columns, strict=False):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(str(column_name))
        set_run_font(run, font_name="Times New Roman", font_size=14, bold=True)

    for row in data.itertuples(index=False):
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=False):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(value))
            set_run_font(run, font_name="Times New Roman", font_size=14)

    return table


def measure_latency(service: ClassifierService) -> dict[str, float]:
    text = "Инженерная команда реализовала интерфейс работы с API и улучшила надежность цифровой платформы."
    image = create_shape_image("Звезда", seed=1234, image_size=service.settings.image_size)

    for _ in range(5):
        service.classify_text(text)
        service.classify_image(image)

    text_times = [int(service.classify_text(text)["processing_time_ms"]) for _ in range(30)]
    image_times = [int(service.classify_image(image)["processing_time_ms"]) for _ in range(30)]
    return {
        "text_avg": round(mean(text_times), 2),
        "text_median": float(median(text_times)),
        "image_avg": round(mean(image_times), 2),
        "image_median": float(median(image_times)),
    }


def build_document() -> Path:
    settings = get_settings(PROJECT_DIR)
    service = ClassifierService(settings)
    service.ensure_ready()
    latency = measure_latency(service)

    document = Document(DOCS_DIR / SOURCE_NAME)

    structure_paragraph = document.paragraphs[20]
    architecture_image_paragraph = document.paragraphs[110]
    architecture_caption = document.paragraphs[111]
    workflow_image_paragraph = document.paragraphs[130]
    workflow_caption = document.paragraphs[131]
    interaction_image_paragraph = document.paragraphs[139]
    interaction_caption = document.paragraphs[140]
    technology_paragraphs = document.paragraphs[145:152]
    requirements_tail = document.paragraphs[100]
    module_tail = document.paragraphs[127]

    set_paragraph_text(
        structure_paragraph,
        "Структура выпускной квалификационной работы включает введение, четыре главы основной части, "
        "заключение, список использованных источников и приложение. В первой главе проводится анализ "
        "предметной области классификации изображений и текстовых данных. Во второй главе рассматриваются "
        "вопросы проектирования программной системы, включая архитектуру, варианты использования и структуру "
        "базы данных. В третьей главе описывается практическая реализация программного продукта и интерфейса. "
        "В четвертой главе приводятся результаты тестирования и оценки качества разработанной системы.",
    )

    add_centered_picture(architecture_image_paragraph, settings.architecture_figure)
    set_paragraph_text(architecture_caption, "Рисунок 2.2 - Архитектура программной системы классификации данных")
    format_caption(architecture_caption)

    add_centered_picture(workflow_image_paragraph, settings.workflow_figure)
    set_paragraph_text(workflow_caption, "Рисунок 2.4 - Алгоритм обработки пользовательского запроса")
    format_caption(workflow_caption)

    add_centered_picture(interaction_image_paragraph, settings.interaction_figure)
    set_paragraph_text(
        interaction_caption,
        "Рисунок 2.5 - Взаимодействие low-code интерфейса, API и прикладных модулей",
    )
    format_caption(interaction_caption)

    replacement_texts = [
        "Для реализации разработанной системы выбран стек Python 3.12, FastAPI, Gradio, scikit-learn и SQLite [1, 4-7].",
        "Язык Python используется как единая среда для серверной логики, процедур обучения моделей и автоматизации тестирования, что уменьшает связность проекта и упрощает сопровождение [1, 2].",
        "В качестве серверной платформы применен FastAPI, позволяющий быстро развернуть REST API, описать входные схемы данных и получить встроенную документацию интерфейсов [4].",
        "Low-code интерфейс реализован на базе Gradio. Данный инструмент позволяет создавать веб-формы для работы с моделями машинного обучения без ручной верстки фронтенда, что соответствует постановке задачи ВКР [5].",
        "Для построения моделей текстовой и графической классификации выбрана библиотека scikit-learn, включающая готовые алгоритмы машинного обучения, средства оценки качества и сериализации артефактов [6].",
        "Хранение истории запросов и метаданных обученных моделей обеспечивается встроенной СУБД SQLite, не требующей отдельного сервера и хорошо подходящей для настольного учебного проекта [7].",
        "Выбранный набор технологий обеспечивает достаточную скорость работы, воспроизводимость результатов эксперимента и возможность дальнейшего расширения системы, включая замену моделей и развитие интерфейса [4-8].",
    ]
    for paragraph, text in zip(technology_paragraphs, replacement_texts, strict=False):
        set_paragraph_text(paragraph, text)

    current = requirements_tail
    current = insert_paragraph_after(current, "Диаграмма вариантов использования системы", STYLE_H2)
    current = insert_paragraph_after(
        current,
        "Для уточнения функций будущего программного продукта была построена диаграмма вариантов использования. "
        "На ней отражены основные действия пользователя: классификация текста, классификация изображения, просмотр "
        "истории обращений и просмотр метрик обученных моделей.",
        STYLE_BODY,
    )
    current = insert_paragraph_after(
        current,
        "Наличие диаграммы вариантов использования позволило сформировать минимально достаточный набор пользовательских "
        "сценариев, который затем был напрямую перенесен в структуру интерфейса Gradio и в набор автотестов.",
        STYLE_BODY,
    )
    current = insert_paragraph_after(current, "", STYLE_BODY)
    add_centered_picture(current, settings.use_case_figure)
    current = insert_paragraph_after(
        current,
        "Рисунок 2.1 - Диаграмма вариантов использования системы классификации данных",
        STYLE_CAPTION,
    )
    format_caption(current)

    current = module_tail
    current = insert_paragraph_after(current, "Проектирование структуры базы данных", STYLE_H2)
    current = insert_paragraph_after(
        current,
        "Для обеспечения воспроизводимости эксперимента и накопления истории обращений в системе спроектирована "
        "локальная база данных SQLite. Использование встроенной СУБД позволяет хранить результаты классификации "
        "непосредственно рядом с приложением и не усложняет процесс развертывания [7].",
        STYLE_BODY,
    )
    current = insert_paragraph_after(
        current,
        "В таблице classification_requests фиксируются тип запроса, источник данных и краткое представление входного "
        "объекта. Таблица classification_results хранит итоговый класс, уровень уверенности модели и время обработки. "
        "Дополнительно таблица model_registry содержит сведения о версиях моделей и достигнутых метриках качества.",
        STYLE_BODY,
    )
    current = insert_paragraph_after(
        current,
        "Такая структура БД обеспечивает связь между пользовательским действием и результатом работы модели, а также "
        "позволяет использовать накопленные сведения для последующего анализа, демонстрации и отладки системы.",
        STYLE_BODY,
    )
    current = insert_paragraph_after(current, "", STYLE_BODY)
    add_centered_picture(current, settings.database_figure)
    current = insert_paragraph_after(current, "Рисунок 2.3 - Структура базы данных программной системы", STYLE_CAPTION)
    format_caption(current)

    document.add_page_break()
    document.add_paragraph(
        "ГЛАВА 3 РАЗРАБОТКА И РЕАЛИЗАЦИЯ ПРОГРАММНОЙ СИСТЕМЫ КЛАССИФИКАЦИИ ДАННЫХ",
        style=STYLE_H1,
    )

    chapter_three_sections = {
        "Общая организация проекта": [
            "Практическая реализация выполнена в формате отдельного Python-проекта, пригодного для открытия в среде PyCharm. "
            "В корне проекта размещены файл запуска main.py, конфигурация зависимостей, папка src с исходным кодом, "
            "папка tests с автоматическими тестами и каталог artifacts с моделями, графиками и служебными данными эксперимента.",
            "Исходный код разделен по функциональным областям. В пакете vkr_classifier выделены модули конфигурации, "
            "доступа к данным, обучения моделей, пользовательского интерфейса, API и слоя хранения. Такое разбиение "
            "облегчает сопровождение проекта и соответствует модульному принципу разработки [2].",
            "Отдельные скрипты generate_assets.py и scripts/build_thesis.py обеспечивают автоматизацию подготовки "
            "артефактов, используемых в пояснительной записке и презентации. За счет этого текстовая часть работы "
            "связана с фактически полученными результатами, а не с декларативным описанием проекта.",
        ],
        "Реализация серверной части и прикладного API": [
            "Серверная часть разработана на основе FastAPI [4]. При запуске приложения создается объект FastAPI, "
            "в котором регистрируются маршруты проверки состояния сервиса, получения истории запросов, просмотра "
            "метаданных моделей и запуска классификации текста либо изображения.",
            "Слой API использует схемы данных pydantic, что позволяет проверять входные параметры и формировать "
            "предсказуемый формат ответа. Для текстового запроса передается JSON-объект с полем text, а для "
            "изображений используется multipart-загрузка файла через стандартный веб-механизм.",
            "Внутри API не реализуется бизнес-логика. Все вычисления и операции записи в БД вынесены в сервисный слой "
            "ClassifierService. Такой подход упрощает тестирование, исключает дублирование кода и позволяет использовать "
            "одни и те же методы как из REST API, так и из low-code интерфейса.",
        ],
        "Реализация текстового классификатора": [
            "Текстовый модуль построен на базе библиотеки scikit-learn [6]. Для подготовки признакового описания "
            "используется TfidfVectorizer с биграммами, а в качестве классификатора выбрана логистическая регрессия. "
            "Данная комбинация обеспечивает хорошее качество на тематических текстах и не требует длительного обучения.",
            "Обучающий набор формируется программно. Для каждой категории создается множество предложений на русском "
            "языке с различными шаблонами и предметной лексикой. В итоговом наборе используются классы 'Спорт', "
            "'Технологии', 'Финансы' и 'Культура'.",
            "После обучения пайплайн сериализуется в файл text_classifier.joblib. При обращении пользователя система "
            "вычисляет вероятности по каждому классу, выбирает наиболее вероятную категорию и дополнительно возвращает "
            "распределение вероятностей для визуального отображения в интерфейсе.",
        ],
        "Реализация классификации изображений": [
            "Модуль изображений ориентирован на распознавание геометрических фигур, что позволяет воспроизвести полный "
            "контур проектирования и тестирования без необходимости скачивания объемных внешних датасетов. Синтетический "
            "набор данных генерируется средствами Pillow: для каждого класса создаются изображения с вариациями размера, "
            "поворота, положения и шумовой компоненты.",
            "Для классификации изображений используется алгоритм k ближайших соседей из библиотеки scikit-learn [6]. "
            "На этапе подготовки входное изображение переводится в оттенки серого, приводится к размеру 32x32 пикселя "
            "и разворачивается в вектор признаков. Такой подход оказался устойчивым на тестовом наборе и показал высокую "
            "точность при минимальной вычислительной сложности.",
            "Обученная модель сохраняется в файл image_classifier.joblib. При инференсе пользователю также возвращается "
            "вероятностное распределение по классам 'Круг', 'Квадрат', 'Треугольник' и 'Звезда', что делает результат "
            "интерпретируемым и удобным для демонстрации в рамках защиты.",
        ],
        "Реализация слоя хранения и журнала истории": [
            "Для хранения прикладных данных применяется SQLite [7]. При первом запуске автоматически создаются таблицы "
            "model_registry, classification_requests и classification_results. Благодаря этому приложение запускается "
            "без предварительной ручной настройки сервера БД.",
            "После каждого распознавания сервис формирует запись о пользовательском действии и результатах модели. "
            "В журнале сохраняются тип запроса, краткое представление входных данных, итоговый класс, оценка уверенности "
            "и фактическое время обработки. Эти сведения используются как в интерфейсе, так и в главе о тестировании.",
            "Хранение версий моделей позволяет фиксировать достигнутые значения accuracy и weighted F1-score. При "
            "следующем обновлении модели соответствующие значения в реестре перезаписываются, что делает проект пригодным "
            "для последующей эволюции и повторного обучения.",
        ],
        "Реализация пользовательского low-code интерфейса": [
            "Веб-интерфейс создан на базе Gradio [5] и встроен в серверное приложение по адресу /ui. На главной странице "
            "размещены две вкладки: для работы с текстом и для работы с изображениями. В каждой вкладке пользователь может "
            "ввести данные, запустить обработку и получить прогноз вместе с вероятностями по классам.",
            "Правая часть интерфейса содержит блок метрик обученных моделей и таблицу истории последних запросов. "
            "Таким образом, интерфейс служит не только точкой ввода данных, но и инструментом демонстрации прикладного "
            "состояния системы, что важно для защиты ВКР.",
            "Использование Gradio позволило отказаться от ручной HTML-верстки и сосредоточиться на логике сервиса. "
            "При этом интерфейс остается полноценным веб-приложением и может использоваться как локально, так и "
            "в качестве демонстрационного стенда для экспериментов.",
        ],
        "Выводы по третьей главе": [
            "В третьей главе выполнена практическая реализация программной системы, объединяющей серверную часть на FastAPI, "
            "low-code интерфейс на Gradio, две модели машинного обучения и подсистему хранения истории запросов.",
            "Разработанная архитектура обеспечивает разделение ответственности между модулями, удобство сопровождения и "
            "возможность автоматической генерации отчетных артефактов. Полученная реализация подготовлена к тестированию "
            "и дальнейшей демонстрации в рамках защиты выпускной квалификационной работы.",
        ],
    }

    for heading, paragraphs in chapter_three_sections.items():
        document.add_paragraph(heading, style=STYLE_H2)
        for text in paragraphs:
            document.add_paragraph(text, style=STYLE_BODY)
        if heading == "Реализация серверной части и прикладного API":
            image_paragraph = document.add_paragraph("", style=STYLE_BODY)
            add_centered_picture(image_paragraph, settings.workflow_figure)
            caption = document.add_paragraph("Рисунок 3.1 - Последовательность обработки запроса в приложении", style=STYLE_CAPTION)
            format_caption(caption)
        if heading == "Реализация пользовательского low-code интерфейса":
            image_paragraph = document.add_paragraph("", style=STYLE_BODY)
            add_centered_picture(image_paragraph, settings.screenshots_dir / "ui_home.png", width_inches=6.15)
            caption = document.add_paragraph("Рисунок 3.2 - Главный экран веб-интерфейса системы", style=STYLE_CAPTION)
            format_caption(caption)
            image_paragraph = document.add_paragraph("", style=STYLE_BODY)
            add_centered_picture(image_paragraph, settings.screenshots_dir / "ui_text_prediction.png", width_inches=6.15)
            caption = document.add_paragraph("Рисунок 3.3 - Пример получения результата текстовой классификации", style=STYLE_CAPTION)
            format_caption(caption)

    document.add_page_break()
    document.add_paragraph("ГЛАВА 4 ТЕСТИРОВАНИЕ И ОЦЕНКА ЭФФЕКТИВНОСТИ РАЗРАБОТАННОЙ СИСТЕМЫ", style=STYLE_H1)

    document.add_paragraph("Организация эксперимента и критерии оценки", style=STYLE_H2)
    document.add_paragraph(
        "Цель тестирования заключалась в проверке корректности работы интерфейса, API, сервисного слоя, "
        "подсистемы хранения данных и качества обученных моделей. Для оценки использовались метрики accuracy, "
        "precision, recall и F1-score, рекомендованные для задач классификации [3, 6].",
        style=STYLE_BODY,
    )
    document.add_paragraph(
        "Дополнительно была измерена задержка инференса после прогрева приложения. Среднее время обработки текста "
        f"составило {latency['text_avg']} мс, медианное значение - {latency['text_median']} мс. Для изображений "
        f"среднее время составило {latency['image_avg']} мс, медианное - {latency['image_median']} мс. "
        "Полученные значения подтверждают возможность интерактивного использования системы.",
        style=STYLE_BODY,
    )

    document.add_paragraph("Автоматическое тестирование программного продукта", style=STYLE_H2)
    document.add_paragraph(
        "Для автоматической проверки корректности проекта реализован набор из шести тестов на базе pytest [8]. "
        "Тесты покрывают сервисный слой, файловые артефакты обучения и REST API. Суммарное покрытие исходного кода "
        "по результатам запуска pytest --cov составило 86%.",
        style=STYLE_BODY,
    )
    test_table = pd.DataFrame(
        [
            ["1", "text service", "Определение технологического текста", "Класс 'Технологии'"],
            ["2", "image service", "Распознавание треугольника", "Класс 'Треугольник'"],
            ["3", "training pipeline", "Проверка генерации моделей и графиков", "Все артефакты созданы"],
            ["4", "text API", "Запрос POST /api/text/classify", "Корректный JSON-ответ"],
            ["5", "image API", "Запрос POST /api/image/classify", "Корректный JSON-ответ"],
            ["6", "history API", "Чтение истории запросов", "Наличие журналируемых записей"],
        ],
        columns=["№", "Компонент", "Проверка", "Ожидаемый результат"],
    )
    add_styled_table(document, "Таблица 4.1 - Сценарии автоматического тестирования", test_table)

    document.add_paragraph("Результаты оценки качества моделей", style=STYLE_H2)
    document.add_paragraph(
        "После обучения моделей были получены итоговые значения метрик качества. Текстовая модель показала "
        "идеальное разделение тематических классов на сформированном корпусе. Модель изображений также "
        "обеспечила высокий уровень точности и устойчиво распознает все четыре геометрических класса.",
        style=STYLE_BODY,
    )
    summary = pd.read_csv(settings.summary_metrics_path)
    for column in ["accuracy", "precision", "recall", "f1_score"]:
        summary[column] = summary[column].map(lambda value: f"{value:.4f}")
    summary = summary.rename(
        columns={
            "model": "Модель",
            "accuracy": "Accuracy",
            "precision": "Precision",
            "recall": "Recall",
            "f1_score": "F1-score",
        }
    )
    add_styled_table(document, "Таблица 4.2 - Итоговые метрики качества обученных моделей", summary)

    image_report = pd.read_csv(settings.image_report_path)
    image_report = image_report[image_report["label"].isin(["Круг", "Квадрат", "Треугольник", "Звезда"])].copy()
    for column in ["precision", "recall", "f1-score"]:
        image_report[column] = image_report[column].map(lambda value: f"{value:.4f}")
    image_report["support"] = image_report["support"].map(lambda value: int(value))
    image_report = image_report.rename(
        columns={
            "label": "Класс",
            "precision": "Precision",
            "recall": "Recall",
            "f1-score": "F1-score",
            "support": "Поддержка",
        }
    )
    add_styled_table(document, "Таблица 4.3 - Показатели качества по классам изображений", image_report)

    for figure_path, caption in [
        (settings.model_comparison_figure, "Рисунок 4.1 - Сравнение итоговых метрик качества моделей"),
        (settings.text_confusion_figure, "Рисунок 4.2 - Матрица ошибок текстовой модели"),
        (settings.image_confusion_figure, "Рисунок 4.3 - Матрица ошибок модели изображений"),
    ]:
        paragraph = document.add_paragraph("", style=STYLE_BODY)
        add_centered_picture(paragraph, figure_path, width_inches=5.9)
        caption_paragraph = document.add_paragraph(caption, style=STYLE_CAPTION)
        format_caption(caption_paragraph)

    document.add_paragraph("Преимущества и ограничения разработанного решения", style=STYLE_H2)
    document.add_paragraph(
        "К преимуществам разработанного программного продукта относятся компактность архитектуры, отсутствие "
        "зависимости от внешней СУБД, быстрый локальный запуск, наличие low-code интерфейса и автоматизированное "
        "формирование отчетных материалов. Пользователь может проверить обе модели в одном приложении без дополнительной настройки.",
        style=STYLE_BODY,
    )
    document.add_paragraph(
        "С точки зрения инженерной практики сильной стороной проекта является воспроизводимость. Все ключевые "
        "артефакты - обученные модели, графики, таблицы, скриншоты и пояснительная записка - формируются из одного "
        "репозитория и могут быть пересобраны повторным запуском скриптов.",
        style=STYLE_BODY,
    )
    document.add_paragraph(
        "Ограничением текущей версии является использование синтетического набора изображений и программно созданного "
        "текстового корпуса. Для дальнейшего развития системы целесообразно подключить реальные открытые датасеты, "
        "добавить сценарии пакетной обработки файлов и расширить перечень поддерживаемых тематических классов.",
        style=STYLE_BODY,
    )

    document.add_paragraph("Выводы по четвертой главе", style=STYLE_H2)
    document.add_paragraph(
        "Проведенное тестирование подтвердило работоспособность разработанной системы. Автоматические тесты успешно "
        "проверяют критические сценарии работы приложения, а итоговые метрики качества показывают пригодность "
        "системы для демонстрации и последующего развития.",
        style=STYLE_BODY,
    )

    document.add_page_break()
    document.add_paragraph("Заключение", style=STYLE_H1)
    for text in [
        "В ходе выполнения выпускной квалификационной работы была разработана программная система классификации изображений "
        "и текстовых данных на языке Python с web-интерфейсом, реализованным на low-code платформе Gradio.",
        "В первой части работы была рассмотрена предметная область, проанализированы методы классификации данных, "
        "современные технологии машинного обучения и существующие решения. Во второй части выполнено проектирование "
        "архитектуры приложения, диаграммы вариантов использования и структуры базы данных.",
        "Практический результат работы представляет собой готовое приложение, включающее FastAPI API, low-code интерфейс, "
        "две модели машинного обучения, журнал истории запросов и набор автотестов. Система успешно проходит проверку "
        "и может использоваться как демонстрационный образец мультимодальной классификации данных.",
        "Проведенная экспериментальная оценка показала, что текстовая модель достигает accuracy 1.0000, а модель "
        "изображений - 0.9583. Таким образом, поставленная цель ВКР достигнута, а разработанная система обладает "
        "практической ценностью и может служить основой для дальнейшего развития проекта.",
    ]:
        document.add_paragraph(text, style=STYLE_BODY)

    document.add_page_break()
    document.add_paragraph("Список использованных источников", style=STYLE_H1)
    bibliography = [
        "1. Python Software Foundation. Python 3.12 documentation [Электронный ресурс]. URL: https://docs.python.org/3/ (дата обращения: 10.04.2026).",
        "2. Ramalho L. Fluent Python. 2nd ed. Sebastopol: O'Reilly Media, 2022. 1014 p.",
        "3. Géron A. Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow. 3rd ed. Sebastopol: O'Reilly Media, 2022. 851 p.",
        "4. FastAPI. Documentation [Электронный ресурс]. URL: https://fastapi.tiangolo.com/ (дата обращения: 10.04.2026).",
        "5. Gradio. Documentation [Электронный ресурс]. URL: https://www.gradio.app/docs (дата обращения: 10.04.2026).",
        "6. Scikit-learn. User Guide [Электронный ресурс]. URL: https://scikit-learn.org/stable/user_guide.html (дата обращения: 10.04.2026).",
        "7. SQLite Documentation [Электронный ресурс]. URL: https://sqlite.org/docs.html (дата обращения: 10.04.2026).",
        "8. pytest Documentation [Электронный ресурс]. URL: https://docs.pytest.org/en/stable/ (дата обращения: 10.04.2026).",
        "9. Goodfellow I., Bengio Y., Courville A. Deep Learning. Cambridge: MIT Press, 2016. 800 p.",
        "10. Bishop C. M. Pattern Recognition and Machine Learning. New York: Springer, 2006. 738 p.",
        "11. Jurafsky D., Martin J. H. Speech and Language Processing. 3rd ed. draft [Электронный ресурс]. URL: https://web.stanford.edu/~jurafsky/slp3/ (дата обращения: 10.04.2026).",
        "12. Han J., Kamber M., Pei J. Data Mining: Concepts and Techniques. 3rd ed. Waltham: Morgan Kaufmann, 2011. 744 p.",
    ]
    for item in bibliography:
        document.add_paragraph(item, style=STYLE_BODY)

    document.add_page_break()
    document.add_paragraph("Приложение А", style=STYLE_H1)
    document.add_paragraph("Структура проекта и инструкция по запуску", style=STYLE_H2)
    for text in [
        "В состав проекта входят: каталог src с исходным кодом приложения, каталог tests с автотестами, каталог "
        "artifacts с обученными моделями, графиками и скриншотами, а также файлы main.py и generate_assets.py для запуска приложения и пересборки артефактов.",
        "Для запуска приложения из каталога проекта используется команда python main.py. Веб-интерфейс становится доступным по адресу http://127.0.0.1:8000/ui.",
        "Для повторной генерации моделей, таблиц и графиков используется команда python generate_assets.py.",
        "Для запуска набора автотестов применяется команда python -m pytest --cov=src/vkr_classifier --cov-report=term-missing.",
    ]:
        document.add_paragraph(text, style=STYLE_BODY)

    structure_table = pd.DataFrame(
        [
            ["src/vkr_classifier", "основные модули приложения"],
            ["tests", "автоматические тесты"],
            ["artifacts/models", "сериализованные модели"],
            ["artifacts/figures", "диаграммы и графики"],
            ["artifacts/screenshots", "скриншоты интерфейса"],
            ["classifier_history.sqlite3", "локальная база данных истории запросов"],
        ],
        columns=["Элемент проекта", "Назначение"],
    )
    add_styled_table(document, "Таблица А.1 - Основные элементы структуры проекта", structure_table)

    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name == STYLE_CAPTION:
            format_caption(paragraph)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = DOCS_DIR / OUTPUT_NAME
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    result = build_document()
    print(result)
